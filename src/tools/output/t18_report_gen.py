"""
T-18 보고서 생성기
분석 결과 전체를 Word(.docx) / CSV 보고서로 자동 생성

라이브러리: python-docx
pyproject.toml 추가: "python-docx>=1.1.0"
"""
import csv
import traceback
from datetime import datetime
from pathlib import Path

from config import SESSION_DIR
from tools.output.t20_trace_logger import log_tool_call


def generate_report(
    session_id: str,
    analysis_result: dict,
    output_format: str = "docx",
) -> dict:
    report_dir = SESSION_DIR / session_id / "reports"
    report_dir.mkdir(parents=True, exist_ok=True)

    timestamp   = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename    = f"report_{timestamp}.{output_format}"
    report_path = report_dir / filename

    try:
        if output_format == "docx":
            _build_docx(report_path, analysis_result)
        else:
            _build_csv(report_path, analysis_result)
        result = {"report_path": str(report_path), "success": True, "error": None}
    except Exception:
        result = {"report_path": "", "success": False, "error": traceback.format_exc()}

    log_tool_call(session_id, "report_gen", {"format": output_format}, result)
    return result


def _build_docx(output_path: Path, data: dict) -> None:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.styles.style import _ParagraphStyle
    except ImportError:
        raise ImportError("python-docx 미설치: uv add python-docx")

    BLUE  = RGBColor(0x2E, 0x74, 0xB5)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    GRAY  = RGBColor(0x70, 0x70, 0x70)

    doc = Document()

    # 페이지 여백
    sec = doc.sections[0]
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(3)
    sec.right_margin  = Cm(2.5)

    # 기본 폰트
    base = doc.styles["Normal"]
    if isinstance(base, _ParagraphStyle):
        base.font.name = "맑은 고딕"
        base.font.size = Pt(10)

    # 제목
    title = doc.add_heading("DAISY 이커머스 분석 보고서", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if title.runs:
        r = title.runs[0]
        r.font.name = "맑은 고딕"
        r.font.size = Pt(18)
        r.font.bold = True
        r.font.color.rgb = BLUE

    date_p = doc.add_paragraph(f"생성일시: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if date_p.runs:
        r = date_p.runs[0]
        r.font.name = "맑은 고딕"
        r.font.size = Pt(9)
        r.font.color.rgb = GRAY
    doc.add_paragraph()

    # 분석 요약
    if summary := data.get("summary"):
        _add_section_heading(doc, "분석 요약", BLUE)
        p = doc.add_paragraph()
        r = p.add_run(str(summary))
        r.font.name = "맑은 고딕"
        r.font.size = Pt(10)
        doc.add_paragraph()

    # 핵심 인사이트
    if insights := data.get("insights"):
        _add_section_heading(doc, "핵심 인사이트", BLUE)
        for i, insight in enumerate(insights, 1):
            p = doc.add_paragraph(style="List Number")
            r = p.add_run(str(insight))
            r.font.name = "맑은 고딕"
            r.font.size = Pt(10)
        doc.add_paragraph()

    # 액션 아이템
    if actions := data.get("actions"):
        _add_section_heading(doc, "액션 아이템", BLUE)
        for action in actions:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(str(action))
            r.font.name = "맑은 고딕"
            r.font.size = Pt(10)
        doc.add_paragraph()

    # KPI 테이블
    if kpi := data.get("kpi_result"):
        _add_section_heading(doc, "KPI 결과", BLUE)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        _set_header_row(table.rows[0], ["지표", "값"], BLUE, WHITE)
        for k, v in kpi.items():
            row = table.add_row()
            val = str(round(v, 4)) if isinstance(v, float) else str(v)
            row.cells[0].text = str(k)
            row.cells[1].text = val
            for cell in row.cells:
                runs = cell.paragraphs[0].runs
                run  = runs[0] if runs else cell.paragraphs[0].add_run(cell.text)
                run.font.name = "맑은 고딕"
                run.font.size = Pt(9)
        _set_col_widths(table, [Cm(8), Cm(5)])
        doc.add_paragraph()

    # 변수 중요도 테이블
    if ranking := data.get("feature_ranking"):
        _add_section_heading(doc, "주요 변수 (상위 5)", BLUE)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        _set_header_row(table.rows[0], ["순위", "변수명", "중요도"], BLUE, WHITE)
        for i, (feat, score) in enumerate(list(ranking.items())[:5], 1):
            row = table.add_row()
            row.cells[0].text = str(i)
            row.cells[1].text = str(feat)
            row.cells[2].text = str(score)
            for cell in row.cells:
                runs = cell.paragraphs[0].runs
                run  = runs[0] if runs else cell.paragraphs[0].add_run(cell.text)
                run.font.name = "맑은 고딕"
                run.font.size = Pt(9)
        _set_col_widths(table, [Cm(2), Cm(7), Cm(4)])
        doc.add_paragraph()

    # 시각화 이미지
    if image_paths := data.get("image_paths"):
        _add_section_heading(doc, "시각화", BLUE)
        for img_path in image_paths:
            p = Path(img_path)
            if p.exists() and p.suffix.lower() == ".png":
                doc.add_picture(str(p), width=Inches(5.5))
                doc.add_paragraph()

    doc.save(str(output_path))


def _add_section_heading(doc, text: str, color) -> None:
    from docx.shared import Pt
    h = doc.add_heading(text, level=1)
    if h.runs:
        r = h.runs[0]
        r.font.name  = "맑은 고딕"
        r.font.size  = Pt(13)
        r.font.bold  = True
        r.font.color.rgb = color


def _set_header_row(row, headers: list, bg_color, text_color) -> None:
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    for cell, header in zip(row.cells, headers):
        cell.text = header
        runs = cell.paragraphs[0].runs
        run  = runs[0] if runs else cell.paragraphs[0].add_run(header)
        run.font.name      = "맑은 고딕"
        run.font.size      = Pt(9)
        run.font.bold      = True
        run.font.color.rgb = text_color
        tc_pr = cell._tc.get_or_add_tcPr()
        shd   = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),
                f"{bg_color[0]:02X}{bg_color[1]:02X}{bg_color[2]:02X}")
        tc_pr.append(shd)


def _set_col_widths(table, widths: list) -> None:
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = width


def _build_csv(output_path: Path, data: dict) -> None:
    with open(output_path, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        writer.writerow(["항목", "내용"])
        writer.writerow(["생성일시", datetime.now().strftime("%Y-%m-%d %H:%M")])
        writer.writerow([])
        if summary := data.get("summary"):
            writer.writerow(["[요약]", summary])
            writer.writerow([])
        if insights := data.get("insights"):
            writer.writerow(["[인사이트]"])
            for i, ins in enumerate(insights, 1):
                writer.writerow([f"{i}.", ins])
            writer.writerow([])
        if actions := data.get("actions"):
            writer.writerow(["[액션 아이템]"])
            for i, act in enumerate(actions, 1):
                writer.writerow([f"{i}.", act])
            writer.writerow([])
        if kpi := data.get("kpi_result"):
            writer.writerow(["[KPI]"])
            writer.writerow(["지표", "값"])
            for k, v in kpi.items():
                writer.writerow([k, v])
            writer.writerow([])
        if ranking := data.get("feature_ranking"):
            writer.writerow(["[변수 중요도]"])
            writer.writerow(["순위", "변수명", "중요도"])
            for i, (feat, score) in enumerate(list(ranking.items())[:5], 1):
                writer.writerow([i, feat, score])